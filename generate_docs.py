"""BookAutoSystem ドキュメント生成スクリプト"""
import os
from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn

OUTPUT_DIR = os.path.join(os.path.dirname(__file__), "data", "output")
os.makedirs(OUTPUT_DIR, exist_ok=True)


def set_cell_shading(cell, color_hex):
    """セルの背景色を設定"""
    from docx.oxml import OxmlElement
    shading_elm = OxmlElement('w:shd')
    shading_elm.set(qn('w:fill'), color_hex)
    shading_elm.set(qn('w:val'), 'clear')
    cell._element.get_or_add_tcPr().append(shading_elm)


def add_table(doc, headers, rows, col_widths=None):
    """表を追加"""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    # ヘッダー
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                run.bold = True
                run.font.size = Pt(9)
        set_cell_shading(cell, '4472C4')
        for run in cell.paragraphs[0].runs:
            run.font.color.rgb = RGBColor(255, 255, 255)
    # データ
    for r_idx, row in enumerate(rows):
        for c_idx, val in enumerate(row):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = str(val)
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(9)
    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Cm(w)
    return table


def setup_doc(title, subtitle="BookAutoSystem"):
    """ドキュメントの初期設定"""
    doc = Document()
    style = doc.styles['Normal']
    style.font.name = 'Yu Gothic'
    style.font.size = Pt(10.5)
    style.element.rPr.rFonts.set(qn('w:eastAsia'), 'Yu Gothic')

    for level in range(1, 4):
        hs = doc.styles[f'Heading {level}']
        hs.font.name = 'Yu Gothic'
        hs.element.rPr.rFonts.set(qn('w:eastAsia'), 'Yu Gothic')

    # 表紙
    for _ in range(6):
        doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(title)
    run.bold = True
    run.font.size = Pt(28)
    run.font.color.rgb = RGBColor(0x2E, 0x4A, 0x7A)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(subtitle)
    run.font.size = Pt(16)
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    for _ in range(4):
        doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("作成日: 2026年4月14日\nバージョン: 1.0")
    run.font.size = Pt(12)

    doc.add_page_break()
    return doc


# ============================================================
# 1. 設計書
# ============================================================
def generate_design_doc():
    doc = setup_doc("システム設計書")

    # 目次
    doc.add_heading('目次', level=1)
    toc_items = [
        "1. システム概要", "2. システム構成", "3. アーキテクチャ設計",
        "4. モジュール設計", "5. データベース設計", "6. 画面設計",
        "7. 外部インターフェース設計", "8. 設定ファイル設計",
        "9. エラーハンドリング設計", "10. セキュリティ設計",
        "11. 非機能要件"
    ]
    for item in toc_items:
        doc.add_paragraph(item, style='List Number')
    doc.add_page_break()

    # 1. システム概要
    doc.add_heading('1. システム概要', level=1)

    doc.add_heading('1.1 目的', level=2)
    doc.add_paragraph(
        'BookAutoSystemは、書店における書籍発注業務を自動化するシステムである。'
        'メールで受信した注文情報を自動的に解析し、複数の仕入先サイトの在庫状況を'
        'Webスクレイピングで確認、最適な仕入先を選定して注文処理を行う。'
        '手作業による発注業務の効率化と、ヒューマンエラーの削減を目的とする。'
    )

    doc.add_heading('1.2 システム機能一覧', level=2)
    add_table(doc,
        ['No.', '機能名', '概要'],
        [
            ['1', 'メール受信・解析', 'Outlook/Graph API/IMAPからメールを取得し、注文情報（注文番号、商品コード、金額等）を抽出'],
            ['2', '自社在庫照合', '自社在庫CSVとの照合により、自社保有の商品を特定'],
            ['3', 'URL指定CSV取込', '外部CSV経由でURLを指定し、特定商品の価格・在庫を取得'],
            ['4', 'Webスクレイピング', '仕入先ECサイト（楽天ブックス、紀伊國屋等）の在庫状況を自動確認'],
            ['5', '仕入先選定', '3段階優先度ロジックにより最適な仕入先を自動選定'],
            ['6', '保留管理', '金額・日数の閾値に基づく保留バケット管理と再スクレイピング'],
            ['7', '自動メール送信', '実店舗仕入先への注文メールを自動生成・送信'],
            ['8', 'Excel出力', '処理台帳・仕入先台帳のExcel出力（色分け付き）'],
            ['9', 'Web管理画面', 'Flask製ダッシュボードでの状態監視・設定変更'],
            ['10', 'GUIランチャー', 'tkinterによる操作パネル（起動/停止/状態表示）'],
            ['11', 'OneDrive連携', 'リモートからの制御指示受信・ステータス同期'],
        ]
    )

    doc.add_heading('1.3 用語定義', level=2)
    add_table(doc,
        ['用語', '説明'],
        [
            ['保留 (HOLD)', '仕入先に割り当て済みだが、金額閾値未到達のため発注保留中の状態'],
            ['保留バケット', '仕入先ごとの保留商品の集約単位。金額・件数・経過日数を追跡'],
            ['閾値到達', '保留バケットの合計金額が仕入先の hold_limit_amount に到達した状態'],
            ['再スクレイピング', '閾値到達後に在庫を再確認する処理。在庫なしの場合はPENDINGに差し戻し'],
            ['処理台帳', '全注文商品の処理状況一覧（Excel出力）'],
            ['仕入先台帳', '仕入先ごとの保留状況一覧（Excel出力）'],
            ['ISBN', '国際標準図書番号（ISBN-10/ISBN-13）'],
            ['ASIN', 'Amazon標準識別番号'],
            ['ECサイト', 'オンライン書店（楽天ブックス、紀伊國屋EC等）'],
            ['実店舗', '物理的な書店（丸善各店舗、紀伊國屋各店舗等）'],
        ]
    )
    doc.add_page_break()

    # 2. システム構成
    doc.add_heading('2. システム構成', level=1)

    doc.add_heading('2.1 技術スタック', level=2)
    add_table(doc,
        ['カテゴリ', '技術/ツール', 'バージョン/備考'],
        [
            ['言語', 'Python', '3.11以上'],
            ['データベース', 'SQLite', 'WALモード、12テーブル'],
            ['メール取得', 'win32com / Microsoft Graph API / IMAP', '3方式切替可能'],
            ['Webスクレイピング', 'Playwright (Chromium)', 'ヘッドレス/有画面切替'],
            ['Web UI', 'Flask + Jinja2', 'ポート5000'],
            ['GUI', 'tkinter', '標準ライブラリ'],
            ['Excel出力', 'openpyxl', '色分け付き台帳'],
            ['HTML解析', 'BeautifulSoup + lxml', 'メール本文解析用'],
            ['配布', 'PyInstaller', 'スタンドアロンexe生成'],
            ['認証', 'MSAL (Azure AD)', 'Graph APIモード時'],
        ]
    )

    doc.add_heading('2.2 ディレクトリ構成', level=2)
    doc.add_paragraph(
        'BookAutoSystem/\n'
        '├── main.py                 # エントリーポイント\n'
        '├── launcher.py             # GUIランチャー\n'
        '├── build.py                # PyInstallerビルド\n'
        '├── requirements.txt        # 依存パッケージ\n'
        '├── config/\n'
        '│   ├── settings.yaml       # システム設定\n'
        '│   ├── suppliers.yaml      # 仕入先マスタ\n'
        '│   ├── mail_patterns.yaml  # メール解析パターン\n'
        '│   └── mail_templates.yaml # メールテンプレート\n'
        '├── src/\n'
        '│   ├── common/             # 共通モジュール（設定,DB,ログ,例外,列挙型）\n'
        '│   ├── mail/               # メール受信（Outlook/Graph/IMAP）\n'
        '│   ├── parser/             # メール解析・コード正規化\n'
        '│   ├── stock/              # 自社在庫・URL指定CSV\n'
        '│   ├── scraper/            # Webスクレイピング（9サイト対応）\n'
        '│   ├── judge/              # 在庫判定・仕入先選定\n'
        '│   ├── hold/               # 保留管理・再スクレイピング\n'
        '│   ├── mailer/             # 自動メール送信\n'
        '│   ├── export/             # Excel出力\n'
        '│   ├── sync/               # OneDrive連携\n'
        '│   └── web/                # Flask Web UI\n'
        '├── data/\n'
        '│   ├── db/                  # SQLiteデータベース\n'
        '│   ├── csv/                 # 入力CSVファイル\n'
        '│   ├── output/              # Excel出力先\n'
        '│   ├── screenshots/         # スクリーンショット\n'
        '│   └── html_snapshots/      # HTMLスナップショット\n'
        '├── logs/                    # ログファイル\n'
        '└── tests/                   # テストコード\n',
        style='No Spacing'
    )
    doc.add_page_break()

    # 3. アーキテクチャ設計
    doc.add_heading('3. アーキテクチャ設計', level=1)

    doc.add_heading('3.1 処理パイプライン', level=2)
    doc.add_paragraph(
        '本システムは10段階のパイプライン処理を polling_interval_sec（デフォルト300秒）間隔で繰り返し実行する。'
        '各ステージは独立しており、1つのステージが失敗しても後続ステージは継続される。'
    )
    add_table(doc,
        ['ステージ', '処理名', '入力', '出力', '担当モジュール'],
        [
            ['1', 'メール取得', 'Outlookメールボックス', 'messagesテーブル (PENDING)', 'mail/outlook_reader.py'],
            ['2', 'メール解析', 'messages (PENDING)', 'order_items (PENDING)', 'parser/mail_parser.py'],
            ['3', '自社在庫CSV取込', 'data/csv/self_stock.xlsx', 'self_stockテーブル', 'stock/self_stock_checker.py'],
            ['4', '自社在庫照合', 'order_items + self_stock', 'order_items (SELF_STOCK)', 'stock/self_stock_checker.py'],
            ['5', 'URL指定CSV取込', 'data/csv/url_specified.xlsx', 'url_csv_entriesテーブル', 'stock/url_csv_loader.py'],
            ['6', 'Webスクレイピング', 'order_items × suppliers', 'scrape_resultsテーブル', 'scraper/orchestrator.py'],
            ['7', '仕入先割当', 'scrape_results', 'order_items (HOLD)', 'judge/supplier_selector.py'],
            ['8', '保留処理', 'order_items (HOLD)', 'hold_buckets/hold_items', 'hold/hold_manager.py'],
            ['9', '再スクレイピング', 'hold_buckets (閾値到達)', 'order_items (ORDERED/PENDING)', 'hold/rescrape_trigger.py'],
            ['10', '自動メール送信', 'order_items (ORDERED)', 'outgoing_mails', 'mailer/auto_mailer.py'],
            ['11', 'Excel出力', 'order_items + suppliers', 'Excel台帳ファイル', 'export/excel_exporter.py'],
        ]
    )

    doc.add_heading('3.2 状態遷移図', level=2)
    doc.add_paragraph('注文商品（order_items）のステータス遷移:')
    doc.add_paragraph(
        'PENDING → SELF_STOCK    （自社在庫に該当）\n'
        'PENDING → HOLD          （仕入先割当成功）\n'
        'PENDING → NO_STOCK      （全仕入先で在庫なし）\n'
        'PENDING → ERROR         （処理エラー）\n'
        'HOLD    → ORDERED       （閾値到達→再スクレイピング→在庫あり）\n'
        'HOLD    → PENDING       （再スクレイピング→在庫なし→差し戻し）\n'
        'ORDERED → CANCELLED     （手動キャンセル）\n'
        '任意    → CANCELLED     （手動キャンセル）',
        style='No Spacing'
    )

    doc.add_heading('3.3 仕入先選定アルゴリズム（3段階優先度）', level=2)
    doc.add_paragraph(
        '仕入先選定は以下の3段階の優先度ロジックで実行される。'
        '在庫ありの仕入先の中から、保留バケットの状態に応じて最適な仕入先を選定する。'
    )
    add_table(doc,
        ['優先度', '条件', '選定基準'],
        [
            ['Tier 1', '保留金額 > 0 かつ < 閾値 (hold_limit_amount)', '保留金額が最小の仕入先を選定（閾値到達を早める）'],
            ['Tier 2', '保留金額 ≧ 閾値', '保留金額が最小の仕入先を選定'],
            ['Tier 3', '保留金額 = 0（新規）', 'priority値の昇順（小さいほど高優先度）'],
        ]
    )
    doc.add_paragraph(
        '※ mail_unit_price_limit を超える高額商品は対象仕入先から除外される。'
    )
    doc.add_page_break()

    # 4. モジュール設計
    doc.add_heading('4. モジュール設計', level=1)

    modules = [
        ('4.1 共通モジュール (src/common/)', [
            ('config.py', 'YAML設定ファイルの読み込み・キャッシュ・書き込み。exe/ソース両方のパス解決。settings.yaml, suppliers.yaml, mail_patterns.yaml, mail_templates.yamlを管理。'),
            ('database.py', 'SQLite接続管理（コンテキストマネージャ）。WALモード、外部キー有効化。12テーブルのCREATE IF NOT EXISTS。マイグレーション対応。'),
            ('logger.py', '3ハンドラ構成のロガー：コンソール、ファイル（日次ローテーション30日）、DB（WARNING以上をlogsテーブルに記録）。'),
            ('exceptions.py', '7種のドメイン固有例外クラス：MailReadError, ParseError, ScrapeError, StockJudgeError, HoldError, MailSendError, ConfigError, DatabaseError。BookAutoError基底クラス。'),
            ('enums.py', '列挙型定義：AvailabilityStatus, OrderStatus, ParseStatus, HoldBucketStatus, SupplierCategory, UrlMode, MailSendStatus。'),
        ]),
        ('4.2 メール受信モジュール (src/mail/)', [
            ('outlook_reader.py', 'win32com.client経由でOutlookアプリケーションに接続。アカウント名・フォルダ名でメールボックスを特定。outlook_message_idで重複排除。メール情報をmessagesテーブルに保存。'),
            ('graph_reader.py', 'Microsoft Graph API v1.0を使用。MSALデバイスコードフローでOAuth2認証。トークンをdata/graph_token_cache.jsonにキャッシュ。日本語フォルダ名を解決。'),
            ('imap_reader.py', 'RFC 3501 IMAP接続。主要プロバイダ自動検出（Outlook.com, Gmail, Yahoo）。Outlook.comはOAuth2 XOAUTH2、他はBasic認証。MIMEヘッダーデコード対応。'),
        ]),
        ('4.3 解析モジュール (src/parser/)', [
            ('mail_parser.py', 'BeautifulSoup+lxmlでHTML→テキスト変換。mail_patterns.yamlの正規表現で注文番号・商品コード・金額等を抽出。1メール→複数商品の分割処理対応。'),
            ('code_normalizer.py', 'ISBN-13/10/ASIN/JANコードの正規化。ISBN-10→13変換（チェックディジット検証）。テキストからの全コード一括抽出。'),
        ]),
        ('4.4 在庫確認モジュール (src/stock/)', [
            ('self_stock_checker.py', 'self_stock.xlsx/.csvの読み込み（ヘッダー名のバリアント検出）。self_stockテーブルを毎回全件置換。order_items(PENDING)との照合でSELF_STOCKステータスに更新。'),
            ('url_csv_loader.py', 'url_specified.xlsx/.tsvからURL指定データを取込。url_csv_entriesテーブルに保存。指定URLをスクレイピングして価格・在庫を取得、最安値にフラグ付与。'),
        ]),
        ('4.5 スクレイピングモジュール (src/scraper/)', [
            ('base_scraper.py', '抽象基底クラス。Playwright Chromiumの起動・終了・ナビゲーション管理。HTMLスナップショット/スクリーンショット保存。テキストノード抽出とパターンマッチングによる在庫判定。ステルスモード対応。'),
            ('orchestrator.py', 'order_items × 有効仕入先のスクレイピング統括。2.5時間以内の結果がある場合はスキップ。エラー時も継続処理。統計情報を返却。'),
            ('url_builder.py', 'URL生成：テンプレートモード（{code}置換）とCSVルックアップモード。suppliers.yamlの設定に基づく。'),
            ('サイト固有スクレイパー (9種)', '楽天ブックス、ヨドバシ、紀伊國屋EC、絵本ナビ、丸善Web、古本屋、丸善店舗、紀伊國屋店舗、三省堂店舗。BaseScraper継承、サイト固有のDOM待機やログイン処理をオーバーライド。'),
        ]),
        ('4.6 判定モジュール (src/judge/)', [
            ('stock_judge.py', '仕入先ごとの最新スクレイピング結果取得。在庫あり仕入先のフィルタリング（priority順）。商品ごとの在庫状況サマリー（available/unavailable/unknown/error件数）。'),
            ('supplier_selector.py', '3段階優先度ロジックによる仕入先選定。order_itemsのplanned_supplier_id更新、ステータスをHOLDに変更。高額商品フィルタリング（mail_unit_price_limit）。'),
        ]),
        ('4.7 保留管理モジュール (src/hold/)', [
            ('hold_manager.py', '保留バケットへの商品追加・バケット作成。バケット再計算（合計金額、件数、最古日付）。閾値判定：hold_limit_amount到達 OR hold_limit_days経過 → THRESHOLD_REACHED。'),
            ('rescrape_trigger.py', '閾値到達バケットの商品を再スクレイピング。在庫あり→保留継続、在庫なし→hold_items削除→order_itemをPENDINGに差し戻し。バケット再計算。'),
        ]),
        ('4.8 メール送信モジュール (src/mailer/)', [
            ('auto_mailer.py', 'ORDERED状態かつauto_mail_enabled仕入先への注文メール送信。mail_templates.yamlのテンプレート使用。Outlook COM/SMTP切替。outgoing_mailsテーブルで送信済み管理（重複防止）。confirmation_mode=trueでログのみ。max_send_per_run=10制限。'),
        ]),
        ('4.9 出力モジュール (src/export/)', [
            ('excel_exporter.py', '処理台帳：全商品×全仕入先マトリックス（色分け：緑=在庫あり、赤=なし、黄=不明、灰=エラー）。仕入先台帳：仕入先ごとの保留状況一覧。ヘッダー固定、オートフィルタ対応。'),
        ]),
        ('4.10 Web UIモジュール (src/web/)', [
            ('app.py', 'Flask Webアプリケーション。ダッシュボード（/）、仕入先一覧（/suppliers）、商品詳細（/item/<id>）、設定（/settings）、ログ（/logs）、パターン編集（/pattern_editor）。スクリーンショット配信。JSON API。'),
        ]),
        ('4.11 同期モジュール (src/sync/)', [
            ('onedrive_sync.py', 'OneDrive経由のリモート制御：停止/一時停止/エクスポート指示の受信。パイプライン結果・システム統計のOneDriveへの書き出し。'),
        ]),
    ]
    for section_title, items in modules:
        doc.add_heading(section_title, level=2)
        add_table(doc,
            ['ファイル', '説明'],
            items
        )
    doc.add_page_break()

    # 5. データベース設計
    doc.add_heading('5. データベース設計', level=1)
    doc.add_paragraph(
        'SQLite 3を使用。WALモード有効化、外部キー制約有効化。'
        'データベースファイル: data/db/main.db'
    )

    doc.add_heading('5.1 テーブル一覧', level=2)
    add_table(doc,
        ['No.', 'テーブル名', '概要', '主なカラム'],
        [
            ['1', 'messages', '受信メール', 'id, outlook_message_id(UNIQUE), account_name, sender, received_at, subject, raw_html, raw_text, parse_status'],
            ['2', 'order_items', '注文商品', 'id, message_id(FK), order_number, product_code_raw, product_code_normalized, product_name, amount, quantity, current_status, planned_supplier_id(FK), assigned_at'],
            ['3', 'self_stock', '自社在庫', 'product_code, stock_qty, imported_at'],
            ['4', 'suppliers', '仕入先マスタ', 'id, supplier_code(UNIQUE), supplier_name, category, priority, hold_limit_amount, hold_limit_days, mail_to_address, scrape_enabled, auto_mail_enabled等'],
            ['5', 'url_rules', 'URL生成ルール', 'supplier_id(FK,UNIQUE), mode, url_template, lookup_csv_path'],
            ['6', 'extraction_patterns', '抽出パターン', 'supplier_id(FK), pattern_name, selector, xpath, class_hint, text_hint, active_flag'],
            ['7', 'scrape_results', 'スクレイピング結果', 'id, order_item_id(FK), supplier_id(FK), scraped_at, availability_status, raw_stock_text, html_snapshot_path, screenshot_path'],
            ['8', 'hold_buckets', '保留バケット', 'id, supplier_id(FK,UNIQUE), total_amount, item_count, oldest_item_date, status'],
            ['9', 'hold_items', '保留明細', 'hold_bucket_id(FK), order_item_id(FK), amount, assigned_at, rescrape_required'],
            ['10', 'outgoing_mails', '送信メール', 'supplier_id(FK), order_item_id(FK), to_address, subject, body, sent_at, send_status'],
            ['11', 'logs', 'システムログ', 'level, module, message, created_at'],
            ['12', 'url_csv_entries', 'URL指定CSV', 'product_code, book_title, page_url, price, is_cheapest, imported_at'],
        ]
    )

    doc.add_heading('5.2 インデックス', level=2)
    add_table(doc,
        ['テーブル', 'カラム', '用途'],
        [
            ['order_items', 'message_id', 'メッセージとの結合'],
            ['order_items', 'current_status', 'ステータス別フィルタリング'],
            ['order_items', 'product_code_normalized', '商品コード検索'],
            ['scrape_results', 'order_item_id', '商品別結果取得'],
            ['scrape_results', 'supplier_id', '仕入先別結果取得'],
            ['hold_items', 'hold_bucket_id', 'バケット別明細取得'],
            ['hold_items', 'order_item_id', '商品別保留状況確認'],
            ['self_stock', 'product_code', '在庫照合'],
        ]
    )

    doc.add_heading('5.3 ER図（概要）', level=2)
    doc.add_paragraph(
        'messages 1---* order_items *---1 suppliers\n'
        'order_items 1---* scrape_results *---1 suppliers\n'
        'order_items 1---* hold_items *---1 hold_buckets 1---1 suppliers\n'
        'order_items 1---* outgoing_mails *---1 suppliers',
        style='No Spacing'
    )
    doc.add_page_break()

    # 6. 画面設計
    doc.add_heading('6. 画面設計', level=1)

    doc.add_heading('6.1 Web UI画面一覧', level=2)
    add_table(doc,
        ['画面名', 'URL', '機能概要'],
        [
            ['ダッシュボード', '/', '注文商品一覧・統計情報（合計/保留中/発注済/エラー件数）・ステータスフィルタ'],
            ['仕入先一覧', '/suppliers', '仕入先台帳・保留バケット状態・金額/件数/経過日数'],
            ['商品詳細', '/item/<id>', 'スクレイピング履歴・スクリーンショット・仕入先別在庫テキスト'],
            ['設定', '/settings', '仕入先の有効/無効・優先度・保留閾値の編集'],
            ['ログ', '/logs', 'システムログ（最新200件）・レベル別フィルタ'],
            ['パターン編集', '/pattern_editor', 'HTML抽出パターンの編集'],
        ]
    )

    doc.add_heading('6.2 GUIランチャー', level=2)
    doc.add_paragraph(
        'tkinterベースの操作パネル。以下の機能を提供:\n'
        '・システム起動/停止ボタン\n'
        '・単発実行（Run Once）ボタン\n'
        '・Web画面を開くボタン\n'
        '・Excel出力ボタン\n'
        '・設定フォルダを開くボタン\n'
        '・リアルタイム統計表示（合計/保留中/発注済/エラー件数）\n'
        '・スクロール可能なログ表示（色分け：INFO=青, WARNING=黄, ERROR=赤, SUCCESS=緑）\n'
        '・ステータスインジケータ（running/stopped/once）\n'
        '・サイクル回数・最終実行時刻・ポーリング間隔の表示'
    )
    doc.add_page_break()

    # 7. 外部インターフェース設計
    doc.add_heading('7. 外部インターフェース設計', level=1)

    doc.add_heading('7.1 入力インターフェース', level=2)
    add_table(doc,
        ['インターフェース', '方式', '形式', '備考'],
        [
            ['注文メール', 'Outlook COM / Graph API / IMAP', 'HTML/テキストメール', '正規表現でフィールド抽出'],
            ['自社在庫', 'ファイル読込', 'Excel (.xlsx) / CSV', 'ISBN列と冊数列を自動検出'],
            ['URL指定CSV', 'ファイル読込', 'Excel (.xlsx) / TSV', '商品コード・URL・価格情報'],
            ['仕入先Webサイト', 'Webスクレイピング (Playwright)', 'HTML', '9サイト対応、ステルスモード'],
            ['OneDrive制御', 'OneDrive API', 'JSON制御ファイル', '停止/一時停止/エクスポート指示'],
        ]
    )

    doc.add_heading('7.2 出力インターフェース', level=2)
    add_table(doc,
        ['インターフェース', '方式', '形式', '備考'],
        [
            ['注文メール', 'Outlook COM / SMTP', 'テキストメール', 'テンプレートベース、confirmation_mode対応'],
            ['処理台帳', 'ファイル出力', 'Excel (.xlsx)', '色分け付き、オートフィルタ対応'],
            ['仕入先台帳', 'ファイル出力', 'Excel (.xlsx)', '保留状況サマリー'],
            ['HTMLスナップショット', 'ファイル保存', 'HTML', '監査証跡用'],
            ['スクリーンショット', 'ファイル保存', 'PNG', '監査証跡用'],
            ['OneDriveステータス', 'OneDrive API', 'JSON', 'サイクル数・統計情報'],
        ]
    )
    doc.add_page_break()

    # 8. 設定ファイル設計
    doc.add_heading('8. 設定ファイル設計', level=1)

    doc.add_heading('8.1 settings.yaml', level=2)
    add_table(doc,
        ['セクション', 'キー', 'デフォルト値', '説明'],
        [
            ['outlook', 'account_name', '(必須)', 'Outlookアカウントのメールアドレスまたは表示名'],
            ['outlook', 'folder_name', '受信トレイ', 'メールフォルダのパス（/区切りで階層指定可）'],
            ['outlook', 'polling_interval_sec', '300', 'ポーリング間隔（秒）'],
            ['outlook', 'mode', 'com', 'メール取得方式: com / graph / imap'],
            ['outlook', 'graph_client_id', '-', 'Azure AD クライアントID（Graphモード時）'],
            ['outlook', 'graph_tenant_id', 'consumers', 'Azure ADテナントID（Graphモード時）'],
            ['database', 'path', 'data/db/main.db', 'SQLiteファイルパス'],
            ['scraping', 'timeout_sec', '30', 'ページ読み込みタイムアウト（秒）'],
            ['scraping', 'headless', 'true', 'ヘッドレスモード'],
            ['scraping', 'max_retries', '2', 'リトライ回数'],
            ['hold', 'strategy', 'fill_smallest', '保留分配戦略'],
            ['mail', 'confirmation_mode', 'false', 'true=ログのみ（送信しない）'],
            ['mail', 'max_send_per_run', '10', '1サイクルあたりの最大送信数'],
            ['web', 'host', '127.0.0.1', 'FlaskバインドIP'],
            ['web', 'port', '5000', 'Flaskバインドポート'],
            ['self_stock', 'csv_path', 'data/csv/self_stock.xlsx', '自社在庫ファイルパス'],
            ['logging', 'dir', 'logs', 'ログ出力ディレクトリ'],
            ['logging', 'level', 'INFO', 'ログレベル'],
        ]
    )

    doc.add_heading('8.2 suppliers.yaml', level=2)
    doc.add_paragraph(
        '仕入先マスタ定義。約40件の仕入先（ECサイト6、実店舗34）を管理。\n'
        '各仕入先エントリの主要フィールド:'
    )
    add_table(doc,
        ['フィールド', '型', '説明'],
        [
            ['code', '文字列', '仕入先コード（一意）'],
            ['name', '文字列', '仕入先名（日本語）'],
            ['category', 'ec/store', 'ECサイト or 実店舗'],
            ['priority', '整数', '優先度（小さいほど高優先度）'],
            ['scrape_enabled', 'boolean', 'スクレイピング有効化'],
            ['auto_mail_enabled', 'boolean', '自動メール送信有効化'],
            ['hold_limit_amount', '整数', '保留閾値金額（円）'],
            ['hold_limit_days', '整数', '保留期限日数'],
            ['mail_unit_price_limit', '整数', '単価上限（超過商品を除外、0=制限なし）'],
            ['mail_to_address', '文字列', 'メール送信先（実店舗のみ）'],
            ['url_mode', 'template/csv_lookup', 'URL生成方式'],
            ['url_template', '文字列', 'URLテンプレート（{code}で置換）'],
            ['positive_patterns', '配列', '在庫あり判定パターン（正規表現）'],
            ['negative_patterns', '配列', '在庫なし判定パターン（正規表現）'],
        ]
    )
    doc.add_page_break()

    # 9. エラーハンドリング設計
    doc.add_heading('9. エラーハンドリング設計', level=1)

    doc.add_heading('9.1 例外クラス階層', level=2)
    doc.add_paragraph(
        'BookAutoError (基底クラス)\n'
        '├── MailReadError      # メール読み取りエラー\n'
        '├── ParseError         # メール解析エラー\n'
        '├── ScrapeError        # スクレイピングエラー\n'
        '├── StockJudgeError    # 在庫判定エラー\n'
        '├── HoldError          # 保留処理エラー\n'
        '├── MailSendError      # メール送信エラー\n'
        '├── ConfigError        # 設定エラー\n'
        '└── DatabaseError      # データベースエラー',
        style='No Spacing'
    )

    doc.add_heading('9.2 エラーハンドリング方針', level=2)
    add_table(doc,
        ['レベル', '方針', '例'],
        [
            ['ステージレベル', '各パイプラインステージをtry/exceptで囲む。エラー時はログ記録し、次ステージに継続', 'メール取得失敗→スクレイピングは実行'],
            ['モジュールレベル', 'ドメイン固有例外をraise。呼び出し元が処理方針を決定', 'ScrapeError → orchestratorがスキップ判定'],
            ['商品レベル', 'スクレイピングループ内で個別商品のエラーをcatch、ログ出力して継続', '1サイトの障害が他サイトに波及しない'],
            ['DBレベル', 'コンテキストマネージャでロールバック制御', '書き込み失敗時は自動ロールバック'],
            ['縮退運転', 'HTMLパターン不一致→UNKNOWN、仕入先未定義→スキップ、メールなし→待機継続', 'システム全体は停止しない'],
        ]
    )

    doc.add_heading('9.3 致命的エラー', level=2)
    doc.add_paragraph(
        '以下のエラーは処理サイクル全体を停止する:\n'
        '・データベース接続失敗\n'
        '・settings.yamlの破損・読み込み不可'
    )
    doc.add_page_break()

    # 10. セキュリティ設計
    doc.add_heading('10. セキュリティ設計', level=1)
    add_table(doc,
        ['項目', '対策'],
        [
            ['認証情報管理', 'Graph API: MSALデバイスコードフロー、トークンファイルは.gitignore。IMAP: パスワードはsettings.yamlに記載（ローカルファイル管理）'],
            ['Webスクレイピング', 'playwright_stealthによるボット検出回避。User-Agent偽装、Accept-Language設定'],
            ['Web UI', 'ローカルバインド（127.0.0.1）のみ。外部アクセス不可'],
            ['SQLインジェクション', 'パラメータ化クエリ（プレースホルダ?使用）。直接文字列結合なし'],
            ['ファイルアクセス', 'data/配下のみにファイル保存。パストラバーサル対策'],
            ['メール送信', 'confirmation_modeによるドライラン対応。max_send_per_runで暴走防止'],
        ]
    )
    doc.add_page_break()

    # 11. 非機能要件
    doc.add_heading('11. 非機能要件', level=1)
    add_table(doc,
        ['項目', '要件', '実装'],
        [
            ['可用性', 'ステージ単位の障害分離', 'try/exceptでステージ独立動作'],
            ['性能', 'ポーリング間隔300秒（調整可能）', 'settings.yaml:polling_interval_sec'],
            ['スケーラビリティ', '1日数百件の注文処理対応', 'SQLite WALモード、非同期スクレイピング'],
            ['保守性', 'モジュール分割設計', '10モジュール、単体テスト完備'],
            ['監査', '全処理の証跡保存', 'DB記録、HTML/スクリーンショット、メール送信ログ'],
            ['配布性', 'スタンドアロン実行', 'PyInstaller exe + Playwright同梱'],
            ['リカバリ', 'ステータス差し戻し', '在庫なし→PENDING差し戻しで自動リトライ'],
        ]
    )

    path = os.path.join(OUTPUT_DIR, "BookAutoSystem_設計書.docx")
    doc.save(path)
    print(f"設計書: {path}")
    return path


# ============================================================
# 2. テスト仕様書
# ============================================================
def generate_test_spec():
    doc = setup_doc("テスト仕様書")

    # 目次
    doc.add_heading('目次', level=1)
    toc = [
        "1. テスト方針", "2. テスト環境", "3. 単体テスト仕様",
        "4. 結合テスト仕様", "5. システムテスト仕様",
        "6. 異常系テスト仕様", "7. 性能テスト仕様", "8. テスト管理"
    ]
    for item in toc:
        doc.add_paragraph(item, style='List Number')
    doc.add_page_break()

    # 1. テスト方針
    doc.add_heading('1. テスト方針', level=1)

    doc.add_heading('1.1 テスト戦略', level=2)
    doc.add_paragraph(
        '本システムのテストは以下の4段階で実施する:\n'
        '1. 単体テスト（Unit Test）: 個別モジュール・関数の動作検証\n'
        '2. 結合テスト（Integration Test）: モジュール間の連携検証\n'
        '3. システムテスト（System Test）: パイプライン全体の動作検証\n'
        '4. 異常系テスト（Error Test）: エラー条件・縮退運転の検証'
    )

    doc.add_heading('1.2 テストツール', level=2)
    add_table(doc,
        ['ツール', '用途'],
        [
            ['pytest', 'テスト実行フレームワーク'],
            ['pytest-cov', 'カバレッジ計測'],
            ['unittest.mock', 'モックオブジェクト生成'],
            ['SQLite (インメモリ)', 'テスト用データベース'],
        ]
    )

    doc.add_heading('1.3 テスト実行方法', level=2)
    doc.add_paragraph(
        'python -m pytest tests/ -v              # 全テスト実行\n'
        'python -m pytest tests/ -v --cov=src    # カバレッジ付き\n'
        'python -m pytest tests/test_xxx.py -v   # 個別テスト実行',
        style='No Spacing'
    )
    doc.add_page_break()

    # 2. テスト環境
    doc.add_heading('2. テスト環境', level=1)
    add_table(doc,
        ['項目', '仕様'],
        [
            ['OS', 'Windows 10 Pro以上'],
            ['Python', '3.11以上'],
            ['データベース', 'SQLite（テスト時はインメモリまたは一時ファイル）'],
            ['ブラウザ', 'Playwright Chromium（スクレイピングテスト時）'],
            ['メール', 'モックオブジェクト（実際のメール接続は不要）'],
            ['ネットワーク', 'スクレイピングテスト時のみ必要'],
        ]
    )
    doc.add_page_break()

    # 3. 単体テスト仕様
    doc.add_heading('3. 単体テスト仕様', level=1)

    # 3.1 コード正規化
    doc.add_heading('3.1 コード正規化テスト (test_code_normalizer.py)', level=2)
    add_table(doc,
        ['No.', 'テストケース', '入力', '期待結果', '検証内容'],
        [
            ['CN-001', 'ISBN-13正規化', '978-4-408-42075-2', '9784408420752', 'ハイフン除去、13桁ISBN'],
            ['CN-002', 'ISBN-10→13変換', '4408420758', '9784408420752', '10桁→13桁変換、チェックディジット'],
            ['CN-003', 'ISBN-10 (Xチェックディジット)', '448084207X', '正規化ISBN-13', 'X末尾の処理'],
            ['CN-004', 'ASIN正規化', 'B08N5WRWNW', 'B08N5WRWNW', 'Bプレフィックス10文字'],
            ['CN-005', 'JANコード', '4988013850019', '4988013850019', '8桁以上数値'],
            ['CN-006', '無効コード', 'ABC123', 'None/空文字列', '正規化不可の場合'],
            ['CN-007', 'テキストからコード抽出', 'ISBN 978-4-408-42075-2 購入', '9784408420752', '文中からのコード検出'],
            ['CN-008', '複数コード抽出', 'ISBN1:978xxx ISBN2:978yyy', '[code1, code2]', '複数コードの一括抽出'],
        ]
    )

    # 3.2 メール解析
    doc.add_heading('3.2 メール解析テスト (test_mail_parser.py)', level=2)
    add_table(doc,
        ['No.', 'テストケース', '入力', '期待結果', '検証内容'],
        [
            ['MP-001', '注文番号抽出', '注文番号：ORD-2026-001', 'ORD-2026-001', '正規表現パターンマッチ'],
            ['MP-002', '商品コード抽出（ISBN-13）', '978-4-408-42075-2', '9784408420752', 'メール本文からISBN抽出'],
            ['MP-003', '金額抽出', '金額：1,500円', '1500', 'カンマ区切り数値の解析'],
            ['MP-004', '数量抽出', '数量：3冊', '3', '数量フィールドの抽出'],
            ['MP-005', 'HTML→テキスト変換', '<p>テスト<br>本文</p>', 'テスト\n本文', 'HTML→プレーンテキスト'],
            ['MP-006', '複数商品メール', '商品1...---...商品2', '2件のorder_items', 'セパレータ分割'],
            ['MP-007', '商品コードなしメール', '注文のご確認', 'parse_status=SKIPPED', 'コードなし→スキップ'],
            ['MP-008', '不正HTML', '<div><span>不正', 'エラーなく解析', 'BeautifulSoupのフォールバック'],
        ]
    )

    # 3.3 在庫チェック
    doc.add_heading('3.3 在庫チェックテスト (test_stock_checker.py)', level=2)
    add_table(doc,
        ['No.', 'テストケース', '入力', '期待結果', '検証内容'],
        [
            ['SC-001', '自社在庫一致', 'order_item(ISBN:978xxx) + self_stock(978xxx:5)', 'SELF_STOCK', '在庫照合成功'],
            ['SC-002', '自社在庫不一致', 'order_item(ISBN:978yyy) + self_stock(978xxx:5)', 'PENDING(変更なし)', '在庫なし→ステータス変更なし'],
            ['SC-003', 'CSV全件置換', '既存self_stock 3件 → 新CSV 2件', 'self_stock 2件', 'テーブル全件置換'],
            ['SC-004', 'CSV形式自動検出', 'ヘッダー「ISBN番号」「在庫数」', '正常読込', 'ヘッダーバリアント対応'],
            ['SC-005', 'CSVファイルなし', 'ファイルパスが存在しない', 'ログ出力、処理継続', 'ファイル不在時の縮退'],
        ]
    )

    # 3.4 保留管理
    doc.add_heading('3.4 保留管理テスト (test_hold_manager.py)', level=2)
    add_table(doc,
        ['No.', 'テストケース', '入力', '期待結果', '検証内容'],
        [
            ['HM-001', 'バケット新規作成', '仕入先Aに初回追加(金額5000)', 'hold_bucket(total=5000,count=1)', 'バケット初期化'],
            ['HM-002', 'バケット追加', '既存バケット(total=5000)に3000追加', 'total=8000, count=2', '金額・件数加算'],
            ['HM-003', '金額閾値到達', 'hold_limit_amount=30000, total=31000', 'status=THRESHOLD_REACHED', '閾値判定'],
            ['HM-004', '日数閾値到達', 'hold_limit_days=14, oldest=15日前', 'status=THRESHOLD_REACHED', '日数超過判定'],
            ['HM-005', '閾値未到達', 'hold_limit_amount=30000, total=20000', 'status=ACTIVE', '閾値未到達'],
            ['HM-006', 'バケット再計算', 'hold_items削除後に再計算', 'total/count更新', '再計算ロジック'],
            ['HM-007', '全バケット再計算', '複数バケット同時再計算', '全バケット正常更新', '一括再計算'],
        ]
    )

    # 3.5 DB
    doc.add_heading('3.5 データベーステスト (test_database.py)', level=2)
    add_table(doc,
        ['No.', 'テストケース', '入力', '期待結果', '検証内容'],
        [
            ['DB-001', 'テーブル作成', 'init_db()呼び出し', '12テーブル作成済み', 'CREATE IF NOT EXISTS'],
            ['DB-002', '接続ライフサイクル', 'with get_connection() as conn:', '正常接続・自動クローズ', 'コンテキストマネージャ'],
            ['DB-003', 'WALモード確認', '接続後のpragma確認', 'journal_mode=wal', 'WAL有効化'],
            ['DB-004', '外部キー有効化', '接続後のpragma確認', 'foreign_keys=1', '外部キー制約'],
            ['DB-005', 'ロールバック', '例外発生時のトランザクション', 'データ未反映', '自動ロールバック'],
        ]
    )

    # 3.6 Enum
    doc.add_heading('3.6 列挙型テスト (test_enums.py)', level=2)
    add_table(doc,
        ['No.', 'テストケース', '入力', '期待結果', '検証内容'],
        [
            ['EN-001', 'OrderStatus値', 'OrderStatus.PENDING', '"PENDING"', '文字列値の確認'],
            ['EN-002', 'AvailabilityStatus値', 'AvailabilityStatus.AVAILABLE', '"AVAILABLE"', '文字列値の確認'],
            ['EN-003', '全ステータス網羅', '各Enumの全メンバー', '期待値一致', '定義の完全性'],
        ]
    )
    doc.add_page_break()

    # 4. 結合テスト
    doc.add_heading('4. 結合テスト仕様', level=1)

    doc.add_heading('4.1 メール取得→解析 結合テスト', level=2)
    add_table(doc,
        ['No.', 'テストケース', '前提条件', '手順', '期待結果'],
        [
            ['IT-001', '正常フロー', 'テストメールデータ準備', '1.fetch_new_mails(mock)\n2.parse_pending_messages()', 'messagesにレコード作成→order_itemsにPENDINGで作成'],
            ['IT-002', '重複メール', '同一メールを2回取得', '1.fetch_new_mails×2', '2回目はoutlook_message_idで重複排除'],
            ['IT-003', '解析エラーメール', '不正形式メール', '1.fetch→2.parse', 'message.parse_status=ERROR, 他は正常処理'],
        ]
    )

    doc.add_heading('4.2 スクレイピング→判定→保留 結合テスト', level=2)
    add_table(doc,
        ['No.', 'テストケース', '前提条件', '手順', '期待結果'],
        [
            ['IT-004', '在庫あり→保留', 'order_item(PENDING), supplier(在庫あり)', '1.scrape→2.select_supplier→3.add_to_hold', 'scrape_results(AVAILABLE)→order_item(HOLD)→hold_items作成'],
            ['IT-005', '全仕入先在庫なし', 'order_item(PENDING), 全supplier(在庫なし)', '1.scrape→2.select_supplier', 'scrape_results(UNAVAILABLE)→order_item(NO_STOCK)'],
            ['IT-006', '閾値到達→再スクレイピング', 'hold_bucket(THRESHOLD_REACHED)', '1.process_rescrape', '在庫あり→ORDERED / 在庫なし→PENDING差し戻し'],
        ]
    )

    doc.add_heading('4.3 保留→メール送信→出力 結合テスト', level=2)
    add_table(doc,
        ['No.', 'テストケース', '前提条件', '手順', '期待結果'],
        [
            ['IT-007', '注文メール送信', 'order_item(ORDERED), auto_mail_enabled', '1.process_auto_mails', 'outgoing_mails(SUCCESS)作成'],
            ['IT-008', '重複送信防止', '既にoutgoing_mails(SUCCESS)あり', '1.process_auto_mails', '送信スキップ'],
            ['IT-009', 'Excel出力', 'order_items複数件あり', '1.export_all', 'Excel台帳ファイル2種生成'],
        ]
    )
    doc.add_page_break()

    # 5. システムテスト
    doc.add_heading('5. システムテスト仕様', level=1)

    doc.add_heading('5.1 パイプライン全体テスト (test_full_pipeline.py)', level=2)
    add_table(doc,
        ['No.', 'テストケース', '前提条件', '手順', '期待結果'],
        [
            ['ST-001', 'フルパイプライン正常系', 'テストデータ一式準備', 'run_pipeline()実行', '全11ステージ正常完了、results["errors"]が空'],
            ['ST-002', '複数サイクル連続実行', 'ST-001完了後', 'run_pipeline()×3回', '重複処理なし、ステータス整合性維持'],
            ['ST-003', 'Webモード起動', '-', 'main.py --web', 'Flask起動、http://127.0.0.1:5000 応答'],
            ['ST-004', '単発実行モード', '-', 'main.py --once', '1サイクルのみ実行して終了'],
            ['ST-005', 'エクスポート専用モード', '-', 'main.py --export', 'Excel台帳のみ出力'],
        ]
    )

    doc.add_heading('5.2 Web UI テスト', level=2)
    add_table(doc,
        ['No.', 'テストケース', '手順', '期待結果'],
        [
            ['ST-006', 'ダッシュボード表示', 'http://127.0.0.1:5000/ アクセス', '注文一覧・統計情報が正しく表示'],
            ['ST-007', '仕入先一覧表示', '/suppliers アクセス', '全仕入先の保留状況が表示'],
            ['ST-008', '商品詳細表示', '/item/1 アクセス', 'スクレイピング履歴・スクリーンショットが表示'],
            ['ST-009', '設定変更', '/settings で優先度変更', 'DB/yamlに変更が反映'],
            ['ST-010', 'ログ表示', '/logs アクセス', '最新200件のログが表示'],
            ['ST-011', 'スクリーンショット配信', '/screenshot/<path>', 'PNG画像がレスポンスとして返却'],
        ]
    )

    doc.add_heading('5.3 GUIランチャーテスト', level=2)
    add_table(doc,
        ['No.', 'テストケース', '手順', '期待結果'],
        [
            ['ST-012', 'システム起動', '「Start System」ボタン押下', 'ステータスがrunningに変化、パイプライン開始'],
            ['ST-013', 'システム停止', '「Stop System」ボタン押下', 'ステータスがstoppedに変化、パイプライン停止'],
            ['ST-014', '単発実行', '「Run Once」ボタン押下', '1サイクルのみ実行、ステータスがonceに変化後stopped'],
            ['ST-015', 'Web画面表示', '「Open Web」ボタン押下', 'デフォルトブラウザでWeb UI表示'],
            ['ST-016', 'リアルタイムログ', 'システム実行中にログ確認', 'ログが色分け表示でリアルタイム更新'],
        ]
    )
    doc.add_page_break()

    # 6. 異常系テスト
    doc.add_heading('6. 異常系テスト仕様', level=1)
    add_table(doc,
        ['No.', 'テストケース', '異常条件', '期待動作'],
        [
            ['ET-001', 'Outlook未起動', 'win32com接続失敗', 'WarningログOut力、メール取得スキップ、他ステージ継続'],
            ['ET-002', 'ネットワーク断', 'スクレイピング時のネットワーク障害', 'ScrapeError、該当仕入先スキップ、他仕入先は継続'],
            ['ET-003', 'サイト構造変更', '在庫判定パターン不一致', 'availability_status=UNKNOWN、ログ警告'],
            ['ET-004', 'DB接続失敗', 'main.dbファイル破損', 'DatabaseError、サイクル全体停止、エラーログ'],
            ['ET-005', 'settings.yaml不正', 'YAML構文エラー', 'ConfigError、システム起動失敗、エラーメッセージ表示'],
            ['ET-006', 'メール送信失敗', 'Outlook送信エラー/SMTP接続不可', 'MailSendError、outgoing_mails(FAILED)記録、他送信は継続'],
            ['ET-007', 'ディスク容量不足', 'スクリーンショット保存失敗', 'WarningログOut力、スクレイピング結果は保存'],
            ['ET-008', 'CSV形式不正', '自社在庫CSVのヘッダー不一致', '警告ログ、在庫チェックスキップ'],
            ['ET-009', 'Graph APIトークン期限切れ', 'OAuth2トークン失効', '再認証フロー起動（デバイスコード表示）'],
            ['ET-010', '同時アクセス', 'DB同時書き込み', 'WALモードにより正常処理'],
        ]
    )
    doc.add_page_break()

    # 7. 性能テスト
    doc.add_heading('7. 性能テスト仕様', level=1)
    add_table(doc,
        ['No.', 'テストケース', '条件', '合格基準'],
        [
            ['PT-001', 'メール解析性能', '100件のメール', '30秒以内に全件解析完了'],
            ['PT-002', 'スクレイピング性能', '50商品×9仕入先', '1サイクル15分以内に完了'],
            ['PT-003', '保留計算性能', '500件のhold_items', 'バケット再計算5秒以内'],
            ['PT-004', 'Excel出力性能', '1000件のorder_items', 'Excel生成30秒以内'],
            ['PT-005', 'Web UIレスポンス', 'ダッシュボード表示', 'レスポンス3秒以内'],
            ['PT-006', 'DB検索性能', 'order_items 10,000件', 'ステータス別検索1秒以内'],
        ]
    )
    doc.add_page_break()

    # 8. テスト管理
    doc.add_heading('8. テスト管理', level=1)

    doc.add_heading('8.1 テスト実施記録テンプレート', level=2)
    add_table(doc,
        ['項目', '内容'],
        [
            ['テスト実施日', 'YYYY/MM/DD'],
            ['テスト実施者', ''],
            ['テスト環境', 'OS / Pythonバージョン'],
            ['テスト対象バージョン', 'コミットハッシュ / タグ'],
            ['テスト結果', 'PASS / FAIL / SKIP'],
            ['不具合件数', ''],
            ['備考', ''],
        ]
    )

    doc.add_heading('8.2 カバレッジ目標', level=2)
    add_table(doc,
        ['モジュール', '目標カバレッジ', '備考'],
        [
            ['src/common/', '90%以上', 'コアロジック'],
            ['src/parser/', '95%以上', '解析精度が業務品質に直結'],
            ['src/judge/', '90%以上', '仕入先選定の正確性'],
            ['src/hold/', '90%以上', '保留管理の正確性'],
            ['src/scraper/', '70%以上', '外部サイト依存部分を除く'],
            ['src/mail/', '70%以上', '外部接続依存部分を除く'],
            ['src/export/', '80%以上', '出力フォーマットの正確性'],
            ['全体', '80%以上', ''],
        ]
    )

    path = os.path.join(OUTPUT_DIR, "BookAutoSystem_テスト仕様書.docx")
    doc.save(path)
    print(f"テスト仕様書: {path}")
    return path


# ============================================================
# 3. 運用マニュアル
# ============================================================
def generate_operation_manual():
    doc = setup_doc("運用マニュアル")

    # 目次
    doc.add_heading('目次', level=1)
    toc = [
        "1. はじめに", "2. システム要件", "3. インストール手順",
        "4. 初期設定", "5. 日常運用", "6. Web管理画面の使い方",
        "7. GUIランチャーの使い方", "8. 設定変更手順",
        "9. トラブルシューティング", "10. メンテナンス",
        "11. バックアップと復旧", "12. FAQ"
    ]
    for item in toc:
        doc.add_paragraph(item, style='List Number')
    doc.add_page_break()

    # 1. はじめに
    doc.add_heading('1. はじめに', level=1)

    doc.add_heading('1.1 本書の目的', level=2)
    doc.add_paragraph(
        '本マニュアルは、BookAutoSystem（書籍発注自動化システム）の運用に必要な'
        '手順・設定・トラブルシューティングをまとめたものです。'
        'システム管理者および日常運用担当者を対象としています。'
    )

    doc.add_heading('1.2 システム概要', level=2)
    doc.add_paragraph(
        'BookAutoSystemは、メールで受信した書籍注文を自動処理するシステムです。\n'
        '主な機能:\n'
        '・メールからの注文情報自動抽出\n'
        '・複数仕入先の在庫自動確認（Webスクレイピング）\n'
        '・最適な仕入先への自動割当・保留管理\n'
        '・実店舗への注文メール自動送信\n'
        '・処理台帳・仕入先台帳のExcel出力\n'
        '・Web画面でのリアルタイム監視'
    )
    doc.add_page_break()

    # 2. システム要件
    doc.add_heading('2. システム要件', level=1)
    add_table(doc,
        ['項目', '要件'],
        [
            ['OS', 'Windows 10 Pro 以上'],
            ['CPU', '2コア以上推奨'],
            ['メモリ', '4GB以上推奨（Playwright使用のため）'],
            ['ストレージ', '500MB以上の空き（スクリーンショット蓄積に応じて増加）'],
            ['ソフトウェア', 'Microsoft Outlook（COMモード時）'],
            ['ネットワーク', 'インターネット接続（スクレイピング・Graph API時）'],
            ['Python', '3.11以上（ソース実行時）'],
            ['ブラウザ', 'Chromium（Playwright同梱、別途インストール不要）'],
        ]
    )
    doc.add_page_break()

    # 3. インストール手順
    doc.add_heading('3. インストール手順', level=1)

    doc.add_heading('3.1 スタンドアロン版（exe）の場合', level=2)
    doc.add_paragraph(
        '1. 配布ZIPファイルを任意のフォルダに展開する\n'
        '   例: C:\\BookAutoSystem\\\n'
        '2. 展開後のフォルダ構成を確認:\n'
        '   BookAutoSystem.exe  - メインプログラム\n'
        '   launcher.py         - GUIランチャー\n'
        '   config/             - 設定ファイル\n'
        '   data/               - データフォルダ\n'
        '   logs/               - ログフォルダ\n'
        '3. config/settings.yaml を編集（「4. 初期設定」参照）\n'
        '4. BookAutoSystem.exe をダブルクリックで起動'
    )

    doc.add_heading('3.2 ソースから実行する場合', level=2)
    doc.add_paragraph(
        '1. Python 3.11以上をインストール\n'
        '2. プロジェクトフォルダでコマンドプロンプトを開く\n'
        '3. 仮想環境を作成・有効化:\n'
        '   python -m venv .venv\n'
        '   .venv\\Scripts\\activate\n'
        '4. 依存パッケージをインストール:\n'
        '   pip install -r requirements.txt\n'
        '5. Playwrightブラウザをインストール:\n'
        '   playwright install chromium\n'
        '6. config/settings.yaml を編集\n'
        '7. 起動:\n'
        '   python main.py',
        style='No Spacing'
    )

    doc.add_heading('3.3 ビルド手順（exe作成）', level=2)
    doc.add_paragraph(
        '1. ソース環境で以下を実行:\n'
        '   python build.py\n'
        '2. dist/BookAutoSystem/ フォルダが生成される\n'
        '3. フォルダをZIP化して配布',
        style='No Spacing'
    )
    doc.add_page_break()

    # 4. 初期設定
    doc.add_heading('4. 初期設定', level=1)

    doc.add_heading('4.1 メール設定 (settings.yaml)', level=2)
    doc.add_paragraph('config/settings.yaml を編集し、メール取得方式を設定する。')

    doc.add_paragraph('【方式1: Outlook COM（推奨）】', style='No Spacing')
    doc.add_paragraph(
        'outlook:\n'
        '  mode: "com"\n'
        '  account_name: "your-email@example.com"\n'
        '  folder_name: "受信トレイ"\n'
        '  polling_interval_sec: 300',
        style='No Spacing'
    )
    doc.add_paragraph('※ Microsoft Outlookがインストール・起動されている必要があります。')

    doc.add_paragraph('【方式2: Microsoft Graph API】', style='No Spacing')
    doc.add_paragraph(
        'outlook:\n'
        '  mode: "graph"\n'
        '  account_name: "your-email@example.com"\n'
        '  graph_client_id: "your-azure-client-id"\n'
        '  graph_tenant_id: "consumers"',
        style='No Spacing'
    )
    doc.add_paragraph('※ Azure AD にアプリケーション登録が必要です。初回起動時にブラウザ認証が表示されます。')

    doc.add_paragraph('【方式3: IMAP】', style='No Spacing')
    doc.add_paragraph(
        'outlook:\n'
        '  mode: "imap"\n'
        '  account_name: "your-email@example.com"\n'
        '  imap_host: "outlook.office365.com"\n'
        '  imap_port: 993\n'
        '  imap_password: "your-password"',
        style='No Spacing'
    )

    doc.add_heading('4.2 仕入先設定 (suppliers.yaml)', level=2)
    doc.add_paragraph(
        'config/suppliers.yaml で仕入先を設定します。\n\n'
        '主要な設定項目:\n'
        '・scrape_enabled: スクレイピングの有効/無効\n'
        '・auto_mail_enabled: 自動メール送信の有効/無効\n'
        '・priority: 優先度（数値が小さいほど高い）\n'
        '・hold_limit_amount: 保留閾値金額（円）\n'
        '・hold_limit_days: 保留期限日数\n'
        '・mail_to_address: メール送信先（実店舗のみ）\n'
        '・positive_patterns: 在庫あり判定パターン\n'
        '・negative_patterns: 在庫なし判定パターン'
    )

    doc.add_heading('4.3 自社在庫ファイルの準備', level=2)
    doc.add_paragraph(
        'data/csv/self_stock.xlsx に自社在庫データを配置します。\n\n'
        '必須列:\n'
        '・ISBN列: 「ISBN」「ISBN番号」「コード」等のヘッダー名（自動検出）\n'
        '・冊数列: 「冊数」「在庫数」「数量」等のヘッダー名（自動検出）\n\n'
        '※ CSV形式 (.csv) も利用可能です。'
    )
    doc.add_page_break()

    # 5. 日常運用
    doc.add_heading('5. 日常運用', level=1)

    doc.add_heading('5.1 起動方法', level=2)
    add_table(doc,
        ['起動方法', 'コマンド/操作', '説明'],
        [
            ['通常起動', 'python main.py', 'パイプライン無限ループ + Web UI + ブラウザ自動起動'],
            ['Web UIのみ', 'python main.py --web', 'Flask Webサーバーのみ起動'],
            ['単発実行', 'python main.py --once', '1サイクルのみ実行して終了'],
            ['Excel出力のみ', 'python main.py --export', 'Excel台帳のみ出力'],
            ['GUIランチャー', 'python launcher.py', 'GUI操作パネルから制御'],
            ['exe版', 'BookAutoSystem.exe', 'スタンドアロン版を起動'],
        ]
    )

    doc.add_heading('5.2 処理サイクルの流れ', level=2)
    doc.add_paragraph(
        '1回のサイクル（デフォルト5分間隔）で以下の処理が順次実行されます:\n\n'
        '① メール取得 → ② メール解析 → ③ 自社在庫CSV取込 → ④ 自社在庫照合\n'
        '→ ⑤ URL指定CSV取込 → ⑥ Webスクレイピング → ⑦ 仕入先割当\n'
        '→ ⑧ 保留処理 → ⑨ 再スクレイピング → ⑩ 自動メール送信 → ⑪ Excel出力\n\n'
        '各ステージは独立しており、1つが失敗しても後続は実行されます。'
    )

    doc.add_heading('5.3 日次確認事項', level=2)
    add_table(doc,
        ['No.', '確認事項', '確認方法', '対応'],
        [
            ['1', 'システム稼働状態', 'Web UI ダッシュボード or GUIランチャー', '停止していれば再起動'],
            ['2', '新規注文の取込状況', 'Web UI ダッシュボード（PENDING件数）', '取り込まれていない場合はメール設定確認'],
            ['3', 'エラー件数', 'Web UI ダッシュボード（ERROR件数）', 'エラー詳細を確認し対処'],
            ['4', '保留バケット状態', 'Web UI 仕入先一覧', '長期保留（>14日）がないか確認'],
            ['5', '自動メール送信状態', 'Web UI ログ', '送信失敗がないか確認'],
            ['6', 'ログファイル', 'logs/app.log', 'WARNING/ERROR がないか確認'],
        ]
    )

    doc.add_heading('5.4 自社在庫ファイルの更新', level=2)
    doc.add_paragraph(
        '自社在庫データを更新する場合:\n'
        '1. data/csv/self_stock.xlsx を最新データで上書き\n'
        '2. 次のサイクルで自動的に取り込まれます\n'
        '   ※ テーブルは毎回全件置換されるため、常に最新データを配置してください'
    )
    doc.add_page_break()

    # 6. Web管理画面
    doc.add_heading('6. Web管理画面の使い方', level=1)
    doc.add_paragraph(
        'ブラウザで http://127.0.0.1:5000/ にアクセスします。\n'
        '※ システム起動時に自動でブラウザが開きます。'
    )

    doc.add_heading('6.1 ダッシュボード（トップページ）', level=2)
    doc.add_paragraph(
        '・注文商品の一覧がテーブル形式で表示されます\n'
        '・画面上部に統計情報が表示されます:\n'
        '  - 合計件数 / PENDING / SELF_STOCK / HOLD / ORDERED / ERROR\n'
        '・ステータスでフィルタリングが可能です\n'
        '・各商品をクリックすると詳細画面に遷移します'
    )

    doc.add_heading('6.2 仕入先一覧', level=2)
    doc.add_paragraph(
        '・全仕入先の保留バケット状態が一覧表示されます\n'
        '・表示項目: 仕入先名、カテゴリ、合計金額、件数、経過日数、閾値、ステータス\n'
        '・THRESHOLD_REACHED のバケットは色付きで強調表示されます'
    )

    doc.add_heading('6.3 商品詳細', level=2)
    doc.add_paragraph(
        '・個別商品のスクレイピング履歴が表示されます\n'
        '・仕入先ごとの最新在庫テキストとスクリーンショットが確認可能\n'
        '・スクリーンショットはクリックで拡大表示できます'
    )

    doc.add_heading('6.4 設定画面', level=2)
    doc.add_paragraph(
        '・仕入先のスクレイピング有効/無効、自動メール有効/無効を切り替えられます\n'
        '・優先度、保留閾値金額、保留期限日数を変更できます\n'
        '・変更は即座にDBとsuppliers.yamlの両方に反映されます'
    )
    doc.add_page_break()

    # 7. GUIランチャー
    doc.add_heading('7. GUIランチャーの使い方', level=1)
    doc.add_paragraph(
        'python launcher.py で起動します。'
    )

    doc.add_heading('7.1 ボタン操作', level=2)
    add_table(doc,
        ['ボタン', '機能', '備考'],
        [
            ['Start System', 'パイプライン無限ループ開始', 'ステータスが「running」に変化'],
            ['Stop System', 'パイプライン停止', '現在のサイクル完了後に停止'],
            ['Run Once', '1サイクルのみ実行', '完了後にステータスが「stopped」に戻る'],
            ['Open Web', 'ブラウザでWeb UIを開く', 'デフォルトブラウザで http://127.0.0.1:5000/'],
            ['Excel Export', 'Excel台帳を即座に出力', 'data/output/ に出力'],
            ['Settings Folder', '設定フォルダを開く', 'config/ フォルダをエクスプローラーで開く'],
        ]
    )

    doc.add_heading('7.2 ステータス表示', level=2)
    doc.add_paragraph(
        '・ステータスインジケータ: running（緑）/ stopped（赤）/ once（青）\n'
        '・サイクル回数: 起動後の累計サイクル数\n'
        '・最終実行時刻: 最後にパイプラインが完了した日時\n'
        '・ポーリング間隔: 現在の設定値\n'
        '・統計情報: 合計/保留中/発注済/エラー件数'
    )
    doc.add_page_break()

    # 8. 設定変更手順
    doc.add_heading('8. 設定変更手順', level=1)

    doc.add_heading('8.1 ポーリング間隔の変更', level=2)
    doc.add_paragraph(
        '1. config/settings.yaml を開く\n'
        '2. outlook.polling_interval_sec の値を変更\n'
        '   例: 600 （10分間隔）\n'
        '3. システムを再起動'
    )

    doc.add_heading('8.2 仕入先の追加', level=2)
    doc.add_paragraph(
        '1. config/suppliers.yaml にエントリを追加\n'
        '2. サイト固有スクレイパーが必要な場合は src/scraper/ にクラスを作成\n'
        '3. positive_patterns / negative_patterns を設定\n'
        '4. システムを再起動（DB同期が自動実行される）'
    )

    doc.add_heading('8.3 メールテンプレートの変更', level=2)
    doc.add_paragraph(
        '1. config/mail_templates.yaml を開く\n'
        '2. store_order の subject / body を編集\n'
        '   使用可能な変数: {supplier_name}, {product_name}, {product_code}, {quantity}, {amount}\n'
        '3. 変更は次のメール送信時から反映されます'
    )

    doc.add_heading('8.4 在庫判定パターンの変更', level=2)
    doc.add_paragraph(
        '仕入先サイトの表示が変更された場合:\n'
        '1. config/suppliers.yaml の該当仕入先の positive_patterns / negative_patterns を修正\n'
        '   例:\n'
        '   positive_patterns:\n'
        '     - "在庫あり"\n'
        '     - "○"\n'
        '     - "取寄可能"\n'
        '   negative_patterns:\n'
        '     - "品切れ"\n'
        '     - "×"\n'
        '     - "販売終了"\n'
        '2. Web UIの「パターン編集」画面からも変更可能です'
    )

    doc.add_heading('8.5 確認モード（ドライラン）', level=2)
    doc.add_paragraph(
        'メール送信を一時停止したい場合:\n'
        '1. config/settings.yaml を開く\n'
        '2. mail.confirmation_mode を true に設定\n'
        '   → メール送信は行われず、ログのみ記録されます\n'
        '3. 復帰時は false に戻してシステムを再起動'
    )
    doc.add_page_break()

    # 9. トラブルシューティング
    doc.add_heading('9. トラブルシューティング', level=1)

    add_table(doc,
        ['症状', '原因', '対処方法'],
        [
            ['メールが取り込まれない', 'Outlookが起動していない', 'Outlookを起動してからシステムを再起動'],
            ['メールが取り込まれない', 'アカウント名/フォルダ名の設定ミス', 'settings.yamlのaccount_name/folder_nameを確認'],
            ['メールが取り込まれない', 'Graph APIトークン切れ', 'data/graph_token_cache.jsonを削除して再認証'],
            ['スクレイピング結果がUNKNOWN', '仕入先サイトの表示変更', 'positive_patterns/negative_patternsを更新'],
            ['スクレイピングがエラー', 'サイトがメンテナンス中', '時間をおいて自動リトライ（次サイクル）'],
            ['スクレイピングがエラー', 'Playwrightブラウザ未インストール', 'playwright install chromium を実行'],
            ['保留が解消されない', '閾値金額が高すぎる', 'suppliers.yamlのhold_limit_amountを下げる'],
            ['保留が長期化', '在庫切れ商品が多い', 'Web UIで商品詳細を確認し手動対応'],
            ['メール送信されない', 'confirmation_modeがtrue', 'settings.yamlでfalseに変更'],
            ['メール送信されない', 'auto_mail_enabledがfalse', 'suppliers.yamlまたはWeb UI設定で有効化'],
            ['メール送信されない', 'mail_to_addressが空', 'suppliers.yamlに送信先アドレスを設定'],
            ['Excel出力されない', 'data/output/フォルダなし', 'data/output/ フォルダを手動作成'],
            ['Web UIが表示されない', 'ポート5000が使用中', 'settings.yamlのweb.portを変更'],
            ['システムが起動しない', 'settings.yamlの構文エラー', 'YAMLの構文をチェック（インデント、コロン後のスペース）'],
            ['システムが起動しない', 'データベース破損', 'data/db/main.dbを削除して再起動（データは初期化される）'],
        ]
    )
    doc.add_page_break()

    # 10. メンテナンス
    doc.add_heading('10. メンテナンス', level=1)

    doc.add_heading('10.1 ログローテーション', level=2)
    doc.add_paragraph(
        '・ログファイルは日次で自動ローテーションされます（30日保持）\n'
        '・ログファイル名: logs/app.log.YYYY-MM-DD\n'
        '・DBログ（logsテーブル）はWARNING以上のみ記録されます'
    )

    doc.add_heading('10.2 スクリーンショットの管理', level=2)
    doc.add_paragraph(
        '・スクレイピング毎にスクリーンショット（PNG）とHTMLスナップショットが保存されます\n'
        '・保存先: data/screenshots/, data/html_snapshots/\n'
        '・ディスク容量を圧迫する場合は、古いファイルを手動で削除してください\n'
        '   推奨: 30日以上経過したファイルを定期的に削除'
    )

    doc.add_heading('10.3 データベースの管理', level=2)
    doc.add_paragraph(
        '・データベースファイル: data/db/main.db\n'
        '・WALモードのため、同時に以下のファイルが存在する場合があります:\n'
        '  - main.db-wal (WALファイル)\n'
        '  - main.db-shm (共有メモリファイル)\n'
        '  ※ これらは正常なファイルです。削除しないでください。\n'
        '・DB容量が肥大化した場合は VACUUM を実行:\n'
        '  sqlite3 data/db/main.db "VACUUM;"'
    )

    doc.add_heading('10.4 Windowsタスクスケジューラでの自動起動', level=2)
    doc.add_paragraph(
        'システム起動時に自動実行する場合:\n\n'
        '1. タスクスケジューラを開く（taskschd.msc）\n'
        '2. 「タスクの作成」を選択\n'
        '3. 全般タブ:\n'
        '   - 名前: BookAutoSystem\n'
        '   - 「ユーザーがログオンしているかどうかにかかわらず実行する」を選択\n'
        '4. トリガータブ:\n'
        '   - 「新規」→「スタートアップ時」を選択\n'
        '5. 操作タブ:\n'
        '   - プログラム: C:\\path\\to\\.venv\\Scripts\\python.exe\n'
        '   - 引数: main.py\n'
        '   - 開始（オプション）: C:\\path\\to\\BookAutoSystem\\\n'
        '6. 条件タブ:\n'
        '   - 「コンピューターをAC電源で使用している場合のみ」のチェックを外す'
    )
    doc.add_page_break()

    # 11. バックアップと復旧
    doc.add_heading('11. バックアップと復旧', level=1)

    doc.add_heading('11.1 バックアップ対象', level=2)
    add_table(doc,
        ['対象', 'ファイル/フォルダ', '重要度', '備考'],
        [
            ['設定ファイル', 'config/', '高', 'カスタマイズ設定が含まれる'],
            ['データベース', 'data/db/main.db', '高', '全処理データが含まれる'],
            ['自社在庫CSV', 'data/csv/', '中', '外部システムから再取得可能な場合もある'],
            ['Excel出力', 'data/output/', '中', '再生成可能'],
            ['スクリーンショット', 'data/screenshots/', '低', '監査用、再生成不可'],
            ['HTMLスナップショット', 'data/html_snapshots/', '低', '監査用、再生成不可'],
            ['Graph APIトークン', 'data/graph_token_cache.json', '低', '再認証で再取得可能'],
        ]
    )

    doc.add_heading('11.2 バックアップ手順', level=2)
    doc.add_paragraph(
        '1. システムを停止する\n'
        '2. 以下のフォルダをバックアップ先にコピー:\n'
        '   - config/ （必須）\n'
        '   - data/db/ （必須）\n'
        '   - data/csv/ （推奨）\n'
        '3. システムを再開する\n\n'
        '※ システム実行中のバックアップはWALモードのため可能ですが、\n'
        '  main.db-wal と main.db-shm も一緒にコピーしてください。'
    )

    doc.add_heading('11.3 復旧手順', level=2)
    doc.add_paragraph(
        '1. システムを停止する\n'
        '2. バックアップファイルを元の場所に復元:\n'
        '   - config/ → プロジェクトルート/config/\n'
        '   - data/db/main.db → data/db/main.db\n'
        '3. システムを起動する\n'
        '4. Web UIでデータの整合性を確認する'
    )

    doc.add_heading('11.4 データベース初期化（リセット）', level=2)
    doc.add_paragraph(
        '全データをリセットして再開する場合:\n'
        '1. システムを停止する\n'
        '2. data/db/main.db を削除する\n'
        '   ※ main.db-wal, main.db-shm があれば同時に削除\n'
        '3. システムを起動する\n'
        '   → テーブルが自動再作成されます\n'
        '   → suppliers.yamlの内容がDBに同期されます\n\n'
        '【注意】この操作により全ての処理履歴が失われます。'
    )
    doc.add_page_break()

    # 12. FAQ
    doc.add_heading('12. FAQ', level=1)

    faqs = [
        ('Q: ポーリング間隔を変更するには？',
         'A: config/settings.yaml の outlook.polling_interval_sec を変更し、システムを再起動してください。'),
        ('Q: 特定の仕入先を一時的に無効にするには？',
         'A: Web UI の設定画面で該当仕入先の scrape_enabled を無効にしてください。または suppliers.yaml で scrape_enabled: false に変更してください。'),
        ('Q: メールの送信テストをしたい',
         'A: settings.yaml で mail.confirmation_mode: true に設定すると、実際の送信は行われずログのみ記録されます。'),
        ('Q: スクレイピングを特定のサイトだけ実行したい',
         'A: suppliers.yaml で対象サイト以外の scrape_enabled を false にしてください。'),
        ('Q: 処理済みの商品を再度PENDINGに戻すには？',
         'A: 現時点ではDBを直接操作する必要があります。SQLiteブラウザ等で order_items テーブルの current_status を PENDING に変更してください。'),
        ('Q: 新しい仕入先サイトに対応するには？',
         'A: 1) suppliers.yaml にエントリ追加 2) 必要に応じて src/scraper/ にスクレイパークラス追加 3) positive/negative_patterns を設定 4) システム再起動'),
        ('Q: ログレベルを変更するには？',
         'A: settings.yaml の logging.level を変更してください（DEBUG/INFO/WARNING/ERROR）。'),
        ('Q: 複数台で同時実行できますか？',
         'A: SQLiteの制約上、同一DBファイルへの複数プロセスからの書き込みは推奨しません。OneDrive連携機能によるリモート制御をご利用ください。'),
    ]
    for q, a in faqs:
        doc.add_paragraph(q, style='List Bullet')
        doc.add_paragraph(a)

    path = os.path.join(OUTPUT_DIR, "BookAutoSystem_運用マニュアル.docx")
    doc.save(path)
    print(f"運用マニュアル: {path}")
    return path


if __name__ == '__main__':
    generate_design_doc()
    generate_test_spec()
    generate_operation_manual()
    print("\n全文書の生成が完了しました。")
